import os
from docx import Document
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate

class DocxTools:
    @staticmethod
    def load_document(filepath):
        """Loads an existing .docx document."""
        return Document(filepath)
    
    @staticmethod
    def save_document(doc, filepath):
        """Saves a document to the specified filepath."""
        doc.save(filepath)

    @staticmethod
    def extract_structure(doc):
        """Extracts text from paragraphs and tables."""
        structure = {
            'paragraphs': [p.text for p in doc.paragraphs],
            'tables': []
        }
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            structure['tables'].append(table_data)
        return structure

    @staticmethod
    def replace_text(doc, old_text, new_text):
        """Safely replaces text while trying to preserve formatting."""
        # Replace in paragraphs
        for p in doc.paragraphs:
            if old_text in p.text:
                for run in p.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            for run in p.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

    @staticmethod
    def add_heading(doc, text, level=1):
        """Adds a heading with the specified level."""
        return doc.add_heading(text, level=level)

    @staticmethod
    def add_paragraph(doc, text, style=None):
        """Adds a paragraph, optionally with a specific style."""
        return doc.add_paragraph(text, style=style)

    @staticmethod
    def add_table(doc, rows, cols):
        """Adds a table with specified rows and columns."""
        return doc.add_table(rows=rows, cols=cols)

    @staticmethod
    def add_image(doc, image_path, width_inches=None):
        """Adds an image to the document."""
        if width_inches:
            return doc.add_picture(image_path, width=Inches(width_inches))
        else:
            return doc.add_picture(image_path)

    @staticmethod
    def create_new_document():
        """Creates a new, empty Document from scratch."""
        return Document()

    @staticmethod
    def render_template(template_path, context, output_path):
        """Renders a Jinja2 template .docx and saves the result."""
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)
