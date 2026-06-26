from docx import Document

def modify_cahier_des_charges(input_path, output_path):
    doc = Document(input_path)
    
    # 1. Update Tech Stack Table
    for table in doc.tables:
        if len(table.rows) > 0 and 'Couche' in table.cell(0, 0).text:
            for row in table.rows:
                couche = row.cells[0].text.strip()
                if couche == 'Frontend':
                    row.cells[1].text = 'React 18 + TypeScript + Vite + Zustand + React Query'
                elif couche == 'Backend':
                    row.cells[1].text = 'Node.js + Express REST API + Zod'
                elif couche == 'Base de données':
                    row.cells[1].text = 'PostgreSQL + Prisma ORM'
                elif couche == 'Authentification':
                    row.cells[1].text = 'JWT (JSON Web Tokens) avec bcryptjs'
            
            # Add an AI row
            new_row = table.add_row()
            new_row.cells[0].text = 'Intelligence Artificielle'
            new_row.cells[1].text = 'Google Generative AI (Gemini)'
            break

    # 2. Update Database Tables Table
    for table in doc.tables:
        if len(table.rows) > 0 and 'Table' in table.cell(0, 0).text and 'Champs principaux' in table.cell(0, 1).text:
            # Add refresh_tokens
            row1 = table.add_row()
            row1.cells[0].text = 'refresh_tokens'
            row1.cells[1].text = 'id, token, userId, expiresAt, createdAt'
            row1.cells[2].text = 'N:1 → users'
            
            # Add system_settings
            row2 = table.add_row()
            row2.cells[0].text = 'system_settings'
            row2.cells[1].text = 'key, value, updatedAt'
            row2.cells[2].text = 'Aucune'
            break

    # Save the updated document
    doc.save(output_path)

if __name__ == '__main__':
    modify_cahier_des_charges(
        r"D:\stage\Projet\Cahier_des_charges_GFB_Hammemi_v3(1).docx", 
        r"D:\stage\Projet\Cahier_des_charges_GFB_Hammemi_v3(1)_updated.docx"
    )
    print("Document successfully updated.")
