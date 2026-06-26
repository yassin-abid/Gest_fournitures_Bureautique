from docx import Document

def extract_docx_content(filepath, output_path):
    doc = Document(filepath)
    with open(output_path, 'w', encoding='utf-8') as f:
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(p.text + '\n')
        
        f.write("\n--- TABLES ---\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for row in table.rows:
                f.write(" | ".join([cell.text.strip() for cell in row.cells]) + '\n')

extract_docx_content(r"D:\stage\Projet\Cahier_des_charges_GFB_Hammemi_v3(1).docx", "scratch_read_docx.txt")
